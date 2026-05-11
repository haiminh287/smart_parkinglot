"""Sinh file Word IEEE references — 10 tài liệu cốt lõi của ParkSmart."""
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Cm, RGBColor

OUT = Path(__file__).parent.parent / "TaiLieuThamKhao.docx"


def add_hyperlink(paragraph, url, text):
    """Thêm hyperlink vào paragraph."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    # Color blue
    c = OxmlElement("w:color")
    c.set(qn("w:val"), "0563C1")
    rPr.append(c)
    # Underline
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_run(paragraph, text, italic=False, bold=False):
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    if italic:
        run.italic = True
    if bold:
        run.bold = True
    return run


def add_reference(doc, num, parts):
    """parts: list of (text, italic_bool) hoặc ('URL', url) cho hyperlink."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.first_line_indent = Cm(0)
    pf.left_indent = Cm(1.0)
    pf.space_after = Pt(6)

    # Số [N]  — hanging indent
    p.paragraph_format.first_line_indent = Cm(-1.0)

    add_run(p, f"[{num}]\t", bold=False)
    for part in parts:
        if isinstance(part, tuple) and part[0] == "URL":
            add_hyperlink(p, part[1], part[1])
        elif isinstance(part, tuple):
            text, italic = part
            add_run(p, text, italic=italic)
        else:
            add_run(p, part)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.bold = True
    if level == 1:
        run.font.size = Pt(14)
    else:
        run.font.size = Pt(12)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(8)


# ─────────────────────────────────────────────────
doc = Document()

# Style mặc định
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)

# Set page margins (A4 chuẩn)
for section in doc.sections:
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

add_heading(doc, "TÀI LIỆU THAM KHẢO", level=1)
doc.add_paragraph()

# ─────────────────────────────────────────────────
# 10 References
# ─────────────────────────────────────────────────

add_reference(doc, 1, [
    "A. Vaswani, N. Shazeer, N. Parmar, J. Uszkoreit, L. Jones, A. N. Gomez, Ł. Kaiser, and I. Polosukhin, ",
    ('"Attention is all you need," in ', False),
    ("Proc. 31st Conf. Neural Inf. Process. Syst. (NIPS)", True),
    (", Long Beach, CA, USA, Dec. 2017, pp. 5998–6008. Available: ", False),
    ("URL", "https://papers.nips.cc/paper/2017/hash/3f5ee243547dee91fbd053c1c4a845aa-Abstract.html"),
])

add_reference(doc, 2, [
    "J. Redmon, S. Divvala, R. Girshick, and A. Farhadi, ",
    ('"You only look once: Unified, real-time object detection," in ', False),
    ("Proc. IEEE Conf. Comput. Vis. Pattern Recognit. (CVPR)", True),
    (", Las Vegas, NV, USA, Jun. 2016, pp. 779–788, doi: 10.1109/CVPR.2016.91. Available: ", False),
    ("URL", "https://doi.org/10.1109/CVPR.2016.91"),
])

add_reference(doc, 3, [
    "M. Tan and Q. V. Le, ",
    ('"EfficientNetV2: Smaller models and faster training," in ', False),
    ("Proc. 38th Int. Conf. Mach. Learn. (ICML)", True),
    (", vol. 139, Jul. 2021, pp. 10096–10106. Available: ", False),
    ("URL", "https://proceedings.mlr.press/v139/tan21a.html"),
])

add_reference(doc, 4, [
    "M. Li, T. Lv, J. Chen, L. Cui, Y. Lu, D. Florencio, C. Zhang, Z. Li, and F. Wei, ",
    ('"TrOCR: Transformer-based optical character recognition with pre-trained models," in ', False),
    ("Proc. AAAI Conf. Artif. Intell.", True),
    (", vol. 37, no. 11, Washington, DC, USA, Feb. 2023, pp. 13094–13102, doi: 10.1609/aaai.v37i11.26538. Available: ", False),
    ("URL", "https://doi.org/10.1609/aaai.v37i11.26538"),
])

add_reference(doc, 5, [
    "P. Lewis, E. Perez, A. Piktus, F. Petroni, V. Karpukhin, N. Goyal, H. Küttler, M. Lewis, W.-t. Yih, T. Rocktäschel, S. Riedel, and D. Kiela, ",
    ('"Retrieval-augmented generation for knowledge-intensive NLP tasks," in ', False),
    ("Proc. 34th Conf. Neural Inf. Process. Syst. (NeurIPS)", True),
    (", Vancouver, BC, Canada, Dec. 2020, pp. 9459–9474. Available: ", False),
    ("URL", "https://papers.nips.cc/paper/2020/hash/6b493230205f780e1bc26945df7481e5-Abstract.html"),
])

add_reference(doc, 6, [
    ("Information Technology — Automatic Identification and Data Capture Techniques — QR Code Bar Code Symbology Specification", True),
    (", ISO/IEC 18004:2015, International Organization for Standardization, Geneva, Switzerland, Feb. 2015. Available: ", False),
    ("URL", "https://www.iso.org/standard/62021.html"),
])

add_reference(doc, 7, [
    'Django Software Foundation, "Django documentation — Django 5.2 with Django REST framework," 2026. Available: ',
    ("URL", "https://docs.djangoproject.com/en/5.2/"),
    (". Accessed: Apr. 29, 2026.", False),
])

add_reference(doc, 8, [
    'S. Ramírez, "FastAPI — Modern, fast (high-performance) web framework for Python APIs," 2026. Available: ',
    ("URL", "https://fastapi.tiangolo.com/"),
    (". Accessed: Apr. 29, 2026.", False),
])

add_reference(doc, 9, [
    'Meta Platforms Inc., "React — A JavaScript library for building user interfaces," 2026. Available: ',
    ("URL", "https://react.dev/"),
    (". Accessed: Apr. 29, 2026.", False),
])

add_reference(doc, 10, [
    'Google LLC, "Gemini API documentation — Google AI for developers," 2026. Available: ',
    ("URL", "https://ai.google.dev/gemini-api/docs"),
    (". Accessed: Apr. 29, 2026.", False),
])

doc.save(OUT)
print(f"Saved: {OUT}")
print(f"Size: {OUT.stat().st_size} bytes")
